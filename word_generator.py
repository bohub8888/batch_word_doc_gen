
import tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
from docx import Document
import os
import threading
import re

def select_file(entry_widget, file_type):
    """Opens a file dialog and sets the selected file path to the entry widget."""
    filepath = filedialog.askopenfilename(
        title=f"请选择 {file_type} 文件",
        filetypes=[(f"{file_type} Files", f"*.{file_type.lower()}")]
    )
    if filepath:
        entry_widget.config(state='normal')
        entry_widget.delete(0, tk.END)
        entry_widget.insert(0, filepath)
        entry_widget.config(state='readonly')

def docx_replace(doc, data):
    """
    Performs a robust replacement of {{key}} placeholders in the entire document.
    This function handles cases where Word splits placeholders across multiple 'runs'.
    """
    # Create a regex pattern from all keys to find all placeholders at once
    # This finds {{key1}}, {{key2}}, etc.
    pattern = re.compile(r"\{\{(" + "|".join(re.escape(k) for k in data.keys()) + r")\}\}")

    # --- Replace in paragraphs ---
    for p in doc.paragraphs:
        # Check if there's any placeholder in the paragraph's full text
        if re.search(pattern, p.text):
            # To handle split runs, we get the full text, do the replacement,
            # and then write it back, which may cause loss of complex formatting.
            full_text = "".join(run.text for run in p.runs)
            
            # Perform all replacements in the concatenated text
            def replace_match(match):
                key = match.group(1)
                value = data.get(key, match.group(0))
                return str(value) if pd.notna(value) else ""

            new_text = re.sub(pattern, replace_match, full_text)

            # Clear the original paragraph content and add the new text
            for run in p.runs:
                run.text = ""
            p.add_run(new_text)

    # --- Replace in tables ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if re.search(pattern, p.text):
                        full_text = "".join(run.text for run in p.runs)
                        
                        def replace_match(match):
                            key = match.group(1)
                            value = data.get(key, match.group(0))
                            return str(value) if pd.notna(value) else ""

                        new_text = re.sub(pattern, replace_match, full_text)
                        
                        for run in p.runs:
                            run.text = ""
                        p.add_run(new_text)


def generate_documents_task(template_path, excel_path, status_label, generate_button):
    """The core logic for generating documents, designed to be run in a separate thread."""
    if not template_path or not excel_path:
        messagebox.showerror("错误", "请先选择 Word 模板和 Excel 数据文件。")
        status_label.config(text="准备就绪")
        generate_button.config(state='normal')
        return

    try:
        status_label.config(text="正在读取 Excel 数据...")
        df = pd.read_excel(excel_path)
        
        required_columns = ['研发活动编号', '研发活动名称']
        if not all(col in df.columns for col in required_columns):
            messagebox.showerror("Excel 文件错误", f"Excel 文件中必须包含 '{required_columns[0]}' 和 '{required_columns[1]}' 列用于文件命名。")
            status_label.config(text="准备就绪")
            generate_button.config(state='normal')
            return

        output_dir = "批量生成的文档"
        os.makedirs(output_dir, exist_ok=True)
        
        total_files = len(df)
        for index, row in df.iterrows():
            status_label.config(text=f"正在生成文件 {index + 1}/{total_files}...")
            
            doc = Document(template_path)
            row_data = row.to_dict()
            
            # Use the new robust replacement function
            docx_replace(doc, row_data)

            activity_id = row['研发活动编号']
            activity_name = row['研发活动名称']
            new_doc_name = f"{activity_id}-{activity_name}.docx"
            save_path = os.path.join(output_dir, new_doc_name)
            doc.save(save_path)

        messagebox.showinfo("成功", f"成功生成 {total_files} 个 Word 文档！\n文件保存在 '{output_dir}' 目录下。")

    except Exception as e:
        messagebox.showerror("发生错误", f"生成过程中出现错误：\n{e}")
    finally:
        status_label.config(text="准备就绪")
        generate_button.config(state='normal')

def start_generation(template_entry, excel_entry, status_label, generate_button):
    """Starts the document generation process in a new thread to avoid freezing the GUI."""
    generate_button.config(state='disabled')
    status_label.config(text="开始处理...")
    
    template_path = template_entry.get()
    excel_path = excel_entry.get()
    
    thread = threading.Thread(
        target=generate_documents_task,
        args=(template_path, excel_path, status_label, generate_button)
    )
    thread.daemon = True
    thread.start()

def main():
    """Creates and runs the Tkinter GUI."""
    root = tk.Tk()
    root.title("批量 Word 文档生成器 (占位符版)")
    root.geometry("500x250")
    root.resizable(False, False)

    frame = tk.Frame(root, padx=10, pady=10)
    frame.pack(fill=tk.BOTH, expand=True)

    tk.Label(frame, text="Word 模板:").grid(row=0, column=0, sticky=tk.W, pady=5)
    template_entry = tk.Entry(frame, width=50, state='readonly')
    template_entry.grid(row=0, column=1, padx=5)
    tk.Button(frame, text="选择...", command=lambda: select_file(template_entry, "docx")).grid(row=0, column=2)

    tk.Label(frame, text="Excel 数据:").grid(row=1, column=0, sticky=tk.W, pady=5)
    excel_entry = tk.Entry(frame, width=50, state='readonly')
    excel_entry.grid(row=1, column=1, padx=5)
    tk.Button(frame, text="选择...", command=lambda: select_file(excel_entry, "xlsx")).grid(row=1, column=2)

    generate_button = tk.Button(
        frame, 
        text="批量生成", 
        font=("Arial", 12, "bold"),
        bg="#4CAF50", 
        fg="white",
        width=20,
        height=2,
        command=lambda: start_generation(template_entry, excel_entry, status_label, generate_button)
    )
    generate_button.grid(row=2, column=0, columnspan=3, pady=20)

    status_label = tk.Label(frame, text="准备就绪", bd=1, relief=tk.SUNKEN, anchor=tk.W)
    status_label.grid(row=3, column=0, columnspan=3, sticky=tk.W+tk.E)

    root.mainloop()

if __name__ == "__main__":
    main()
